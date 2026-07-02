import argparse
import csv
import json
import re
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any


ARTICLE_NUMBER_PATTERN = re.compile(r"\d+(?:\.\d+)*")
SUCCESS_RELEVANCE = {"primary", "secondary"}
ALLOWED_RELEVANCE = SUCCESS_RELEVANCE | {"acceptable"}

DOCUMENT_ALIASES = {
    "tk_rf": (
        "tk_rf",
        "trudovoy_kodeks_rf",
        "trudovoj-kodeks-rossijskoj-federatsii",
        "трудовой кодекс российской федерации",
    ),
    "uk_rf": (
        "uk_rf",
        "ugolovnyy_kodeks_rf",
        "уголовный кодекс российской федерации",
    ),
    "gk_rf": (
        "gk_rf",
        "grazhdanskiy_kodeks_rf",
        "гражданский кодекс российской федерации",
        "гражданский кодекс российской федерации часть первая",
        "гражданский кодекс российской федерации часть вторая",
    ),
    "nk_rf": (
        "nk_rf",
        "nalogovyy_kodeks_rf",
        "налоговый кодекс российской федерации",
        "налоговый кодекс российской федерации часть первая",
        "налоговый кодекс российской федерации часть вторая",
    ),
    "zkpp_rf": (
        "zkpp_rf",
        "zakon_o_zashite_prav_potrebiteley",
        "закон российской федерации о защите прав потребителей",
        "закон о защите прав потребителей",
        "защите прав потребителей",
    ),
    "sk_rf": (
        "sk_rf",
        "semeynyy_kodeks_rf",
        "семейный кодекс российской федерации",
    ),
    "zhk_rf": (
        "zhk_rf",
        "zhilishchnyy_kodeks_rf",
        "жилищный кодекс российской федерации",
    ),
}


def normalize_article_number(value: object) -> str | None:
    if value is None:
        return None

    match = ARTICLE_NUMBER_PATTERN.search(str(value))
    return match.group(0) if match else None


def normalize_document_identity(value: object) -> str:
    if value is None:
        return ""
    text = str(value).casefold().replace("ё", "е")
    return re.sub(r"[^a-zа-я0-9]+", "", text)


def _alias_key(value: object) -> str | None:
    normalized_value = normalize_document_identity(value)
    if not normalized_value:
        return None

    for document_id, aliases in DOCUMENT_ALIASES.items():
        normalized_id = normalize_document_identity(document_id)
        normalized_aliases = {
            normalize_document_identity(alias)
            for alias in aliases
        }
        if (
            normalized_value == normalized_id
            or normalized_value.startswith(normalized_id)
            or any(
                normalized_value == alias
                or alias in normalized_value
                or normalized_value in alias
                for alias in normalized_aliases
            )
        ):
            return document_id
    return None


def _expected_document_aliases(
    case: dict[str, Any],
    expected_document_id: object | None = None,
) -> set[str]:
    expected_id = expected_document_id or case.get("expected_document_id")
    expected_title = (
        None if expected_document_id else case.get("expected_document_title")
    )
    alias_key = _alias_key(expected_id) if expected_id else _alias_key(expected_title)

    if alias_key:
        return {
            normalize_document_identity(alias)
            for alias in DOCUMENT_ALIASES[alias_key]
        }

    fallback_value = expected_id or expected_title
    normalized_fallback = normalize_document_identity(fallback_value)
    return {normalized_fallback} if normalized_fallback else set()


def _document_matches(
    case: dict[str, Any],
    result: dict[str, Any],
    expected_document_id: object | None = None,
) -> bool:
    expected_aliases = _expected_document_aliases(
        case,
        expected_document_id=expected_document_id,
    )
    if not expected_aliases:
        return True

    result_values = (
        result.get("document_title"),
        result.get("filename"),
        result.get("source_filename"),
        result.get("document_id"),
    )
    normalized_results: set[str] = set()
    for value in result_values:
        normalized_value = normalize_document_identity(value)
        if normalized_value:
            normalized_results.add(normalized_value)

    return any(
        alias in result_value or result_value in alias
        for alias in expected_aliases
        for result_value in normalized_results
    )


def expected_sources_for_case(case: dict[str, Any]) -> list[dict[str, str]]:
    expected_sources = case.get("expected_sources")
    if isinstance(expected_sources, list) and expected_sources:
        normalized_sources = []
        for source in expected_sources:
            if not isinstance(source, dict):
                continue
            article = normalize_article_number(source.get("article_number"))
            document_id = str(source.get("document_id") or "").strip()
            relevance = str(source.get("relevance") or "").strip().casefold()
            if article and document_id and relevance in ALLOWED_RELEVANCE:
                normalized_sources.append(
                    {
                        "document_id": document_id,
                        "article_number": article,
                        "relevance": relevance,
                        "reason": str(source.get("reason") or "").strip(),
                    }
                )
        return normalized_sources

    document_id = str(case.get("expected_document_id") or "").strip()
    return [
        {
            "document_id": document_id,
            "article_number": article,
            "relevance": "primary",
            "reason": "legacy expected_article_numbers",
        }
        for value in case.get("expected_article_numbers", [])
        if (article := normalize_article_number(value))
    ]


def match_result_relevance(
    case: dict[str, Any],
    result: dict[str, Any],
) -> str | None:
    result_article = normalize_article_number(result.get("article_number"))
    if not result_article:
        return None

    matched_relevances = []
    for source in expected_sources_for_case(case):
        if (
            result_article == source["article_number"]
            and _document_matches(
                case,
                result,
                expected_document_id=source["document_id"],
            )
        ):
            matched_relevances.append(source["relevance"])
    for relevance in ("primary", "secondary", "acceptable"):
        if relevance in matched_relevances:
            return relevance
    return None


def is_expected_match(case: dict[str, Any], result: dict[str, Any]) -> bool:
    return match_result_relevance(case, result) in SUCCESS_RELEVANCE


def result_matches_expected_document(
    case: dict[str, Any],
    result: dict[str, Any],
) -> bool:
    sources = expected_sources_for_case(case)
    return any(
        _document_matches(
            case,
            result,
            expected_document_id=source["document_id"],
        )
        for source in sources
        if source["relevance"] in SUCCESS_RELEVANCE
    )


def _optional_float(value: object) -> float | None:
    if value is None:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def evaluate_case_results(
    case: dict[str, Any],
    results: list[dict[str, Any]],
    top_k: int = 20,
) -> dict[str, Any]:
    if case.get("case_type") == "out_of_scope":
        return {
            "case": case,
            "top_results": [],
            "matched_result": None,
            "matched_relevance": None,
            "partial_match": False,
            "not_evaluated": True,
            "hit_top_1": None,
            "hit_top_3": None,
            "hit_top_5": None,
            "hit_top_10": None,
            "hit_top_20": None,
            "rank": None,
            "reciprocal_rank": None,
            "document_hit_top_5": None,
            "wrong_document_top_1": None,
            "best_similarity": None,
            "best_distance": None,
            "semantic_score": None,
            "keyword_score": None,
            "hybrid_score": None,
        }

    rank: int | None = None
    matched_result: dict[str, Any] | None = None
    matched_relevance: str | None = None
    acceptable_result: dict[str, Any] | None = None

    for index, result in enumerate(results, start=1):
        relevance = match_result_relevance(case, result)
        if relevance in SUCCESS_RELEVANCE:
            rank = index
            matched_result = result
            matched_relevance = relevance
            break
        if relevance == "acceptable" and acceptable_result is None:
            acceptable_result = result

    first_result = results[0] if results else {}
    document_hit_top_5 = any(
        result_matches_expected_document(case, result)
        for result in results[:5]
    )
    wrong_document_top_1 = bool(
        first_result
        and not result_matches_expected_document(case, first_result)
    )

    def hit_at(depth: int) -> bool | None:
        if top_k < depth:
            return None
        return rank is not None and rank <= depth

    return {
        "case": case,
        "top_results": results,
        "matched_result": matched_result or acceptable_result,
        "matched_relevance": (
            matched_relevance
            or ("acceptable" if acceptable_result else None)
        ),
        "partial_match": matched_result is None and acceptable_result is not None,
        "not_evaluated": False,
        "hit_top_1": hit_at(1),
        "hit_top_3": hit_at(3),
        "hit_top_5": hit_at(5),
        "hit_top_10": hit_at(10),
        "hit_top_20": hit_at(20),
        "rank": rank,
        "reciprocal_rank": 1.0 / rank if rank else 0.0,
        "document_hit_top_5": (
            document_hit_top_5 if top_k >= 5 else None
        ),
        "wrong_document_top_1": wrong_document_top_1,
        "best_similarity": _optional_float(first_result.get("similarity")),
        "best_distance": _optional_float(first_result.get("distance")),
        "semantic_score": _optional_float(first_result.get("semantic_score")),
        "keyword_score": _optional_float(first_result.get("keyword_score")),
        "hybrid_score": _optional_float(first_result.get("hybrid_score")),
    }


def calculate_summary_metrics(
    evaluations: list[dict[str, Any]],
) -> dict[str, Any]:
    scored_evaluations = [
        item
        for item in evaluations
        if not item.get("not_evaluated")
    ]
    count = len(scored_evaluations)
    similarities = [
        item["best_similarity"]
        for item in scored_evaluations
        if item.get("best_similarity") is not None
    ]
    ranks = [
        item["rank"]
        for item in scored_evaluations
        if item.get("rank") is not None
    ]
    document_hits = [
        item["document_hit_top_5"]
        for item in scored_evaluations
        if item.get("document_hit_top_5") is not None
    ]
    wrong_documents = [
        item["wrong_document_top_1"]
        for item in scored_evaluations
        if item.get("wrong_document_top_1") is not None
    ]

    metrics: dict[str, Any] = {
        "questions_count": count,
        "out_of_scope_count": len(evaluations) - count,
        "mrr": (
            sum(item["reciprocal_rank"] for item in scored_evaluations) / count
            if count
            else 0.0
        ),
        "average_first_similarity": (
            sum(similarities) / len(similarities) if similarities else 0.0
        ),
        "mean_rank": sum(ranks) / len(ranks) if ranks else None,
        "document_recall_at_5": (
            sum(bool(value) for value in document_hits) / len(document_hits)
            if document_hits
            else None
        ),
        "wrong_document_at_1": (
            sum(bool(value) for value in wrong_documents) / len(wrong_documents)
            if wrong_documents
            else None
        ),
    }
    for depth in (1, 3, 5, 10, 20):
        values = [
            item[f"hit_top_{depth}"]
            for item in scored_evaluations
            if item.get(f"hit_top_{depth}") is not None
        ]
        metrics[f"hit_top_{depth}_count"] = (
            sum(bool(value) for value in values) if values else None
        )
        metrics[f"top_{depth}_accuracy"] = (
            sum(bool(value) for value in values) / len(values)
            if values
            else None
        )
    return metrics


def calculate_grouped_metrics(
    evaluations: list[dict[str, Any]],
) -> dict[str, dict[str, Any]]:
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for evaluation in evaluations:
        if evaluation.get("not_evaluated"):
            continue
        grouped[evaluation["case"].get("law") or "Без названия"].append(evaluation)

    return {
        law: calculate_summary_metrics(items)
        for law, items in sorted(grouped.items())
    }


def calculate_question_type_metrics(
    evaluations: list[dict[str, Any]],
) -> dict[str, dict[str, Any]]:
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for evaluation in evaluations:
        if evaluation.get("not_evaluated"):
            continue
        question_type = (
            evaluation["case"].get("question_type") or "unspecified"
        )
        grouped[str(question_type)].append(evaluation)

    return {
        question_type: calculate_summary_metrics(items)
        for question_type, items in sorted(grouped.items())
    }


def load_cases(cases_path: Path) -> list[dict[str, Any]]:
    try:
        data = json.loads(cases_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise RuntimeError(f"Cannot load evaluation cases from {cases_path}: {exc}") from exc

    if not isinstance(data, list) or not data:
        raise RuntimeError("Evaluation cases must be a non-empty JSON array")

    seen_ids: set[str] = set()
    for index, case in enumerate(data, start=1):
        if not isinstance(case, dict):
            raise RuntimeError(f"Evaluation case #{index} must be an object")
        case_id = str(case.get("id") or "").strip()
        if not case_id:
            raise RuntimeError(f"Evaluation case #{index} has no id")
        if case_id in seen_ids:
            raise RuntimeError(f"Duplicate evaluation case id: {case_id}")
        if not str(case.get("question") or "").strip():
            raise RuntimeError(f"Evaluation case {case_id} has no question")
        if case.get("case_type") == "out_of_scope":
            if case.get("expected_behavior") != "no_answer":
                raise RuntimeError(
                    f"Out-of-scope case {case_id} must expect no_answer"
                )
        else:
            raw_sources = case.get("expected_sources")
            if raw_sources is not None:
                if not isinstance(raw_sources, list) or not raw_sources:
                    raise RuntimeError(
                        f"Evaluation case {case_id} has invalid expected_sources"
                    )
                for source in raw_sources:
                    if (
                        not isinstance(source, dict)
                        or not str(source.get("document_id") or "").strip()
                        or not normalize_article_number(
                            source.get("article_number")
                        )
                        or source.get("relevance") not in ALLOWED_RELEVANCE
                    ):
                        raise RuntimeError(
                            f"Evaluation case {case_id} has invalid expected source"
                        )
            if not expected_sources_for_case(case):
                raise RuntimeError(
                    f"Evaluation case {case_id} has no valid expected sources"
                )
        seen_ids.add(case_id)

    return data


def _result_document_name(result: dict[str, Any]) -> str:
    return str(
        result.get("document_title")
        or result.get("filename")
        or result.get("source_filename")
        or result.get("document_id")
        or "Документ не указан"
    )


def _format_similarity(value: object) -> str:
    number = _optional_float(value)
    return f"{number:.4f}" if number is not None else "—"


def _format_top_results(
    results: list[dict[str, Any]],
    limit: int | None = None,
) -> str:
    selected = results[:limit] if limit is not None else results
    lines: list[str] = []
    for index, result in enumerate(selected, start=1):
        article = normalize_article_number(result.get("article_number")) or "—"
        title = result.get("article_title") or "без названия"
        lines.append(
            f"{index}) {_result_document_name(result)}, статья {article}, "
            f"{title}, similarity={_format_similarity(result.get('similarity'))}"
        )
    return " | ".join(lines) if lines else "Результатов нет"


def _format_expected_sources(case: dict[str, Any]) -> str:
    sources = expected_sources_for_case(case)
    return "; ".join(
        (
            f"{source['document_id']} — ст. {source['article_number']} "
            f"({source['relevance']})"
        )
        for source in sources
    ) or "—"


def write_csv_report(
    evaluations: list[dict[str, Any]],
    output_path: Path,
    metadata: dict[str, Any] | None = None,
) -> None:
    metadata = metadata or {}
    fieldnames = [
        "case_id",
        "case_type",
        "question_type",
        "difficulty",
        "expected_behavior",
        "retrieval_mode",
        "hybrid_semantic_weight",
        "hybrid_keyword_weight",
        "hybrid_metadata_weight",
        "law",
        "question",
        "expected_articles",
        "hit_top_1",
        "hit_top_3",
        "hit_top_5",
        "hit_top_10",
        "hit_top_20",
        "rank",
        "reciprocal_rank",
        "matched_relevance",
        "partial_match",
        "document_hit_top_5",
        "wrong_document_top_1",
        "best_similarity",
        "best_distance",
        "semantic_score",
        "keyword_score",
        "hybrid_score",
        "top_results",
    ]
    with output_path.open("w", encoding="utf-8-sig", newline="") as file:
        writer = csv.DictWriter(file, fieldnames=fieldnames)
        writer.writeheader()
        for evaluation in evaluations:
            case = evaluation["case"]
            writer.writerow(
                {
                    "case_id": case["id"],
                    "case_type": case.get("case_type", "retrieval"),
                    "question_type": case.get("question_type"),
                    "difficulty": case.get("difficulty"),
                    "expected_behavior": case.get("expected_behavior"),
                    "retrieval_mode": metadata.get("retrieval_mode"),
                    "hybrid_semantic_weight": (
                        metadata.get("hybrid_semantic_weight")
                        if metadata.get("retrieval_mode") == "hybrid"
                        else None
                    ),
                    "hybrid_keyword_weight": (
                        metadata.get("hybrid_keyword_weight")
                        if metadata.get("retrieval_mode") == "hybrid"
                        else None
                    ),
                    "hybrid_metadata_weight": (
                        metadata.get("hybrid_metadata_weight")
                        if metadata.get("retrieval_mode") == "hybrid"
                        else None
                    ),
                    "law": case.get("law"),
                    "question": case.get("question"),
                    "expected_articles": _format_expected_sources(case),
                    "hit_top_1": evaluation["hit_top_1"],
                    "hit_top_3": evaluation["hit_top_3"],
                    "hit_top_5": evaluation["hit_top_5"],
                    "hit_top_10": evaluation["hit_top_10"],
                    "hit_top_20": evaluation["hit_top_20"],
                    "rank": evaluation["rank"],
                    "reciprocal_rank": evaluation["reciprocal_rank"],
                    "matched_relevance": evaluation["matched_relevance"],
                    "partial_match": evaluation["partial_match"],
                    "document_hit_top_5": evaluation["document_hit_top_5"],
                    "wrong_document_top_1": evaluation["wrong_document_top_1"],
                    "best_similarity": evaluation["best_similarity"],
                    "best_distance": evaluation["best_distance"],
                    "semantic_score": evaluation["semantic_score"],
                    "keyword_score": evaluation["keyword_score"],
                    "hybrid_score": evaluation["hybrid_score"],
                    "top_results": _format_top_results(evaluation["top_results"]),
                }
            )


def write_json_report(
    evaluations: list[dict[str, Any]],
    summary: dict[str, Any],
    grouped_metrics: dict[str, dict[str, Any]],
    metadata: dict[str, Any],
    output_path: Path,
    question_type_metrics: dict[str, dict[str, Any]] | None = None,
) -> None:
    payload = {
        **metadata,
        "summary": summary,
        "metrics_by_law": grouped_metrics,
        "metrics_by_question_type": question_type_metrics or {},
        "cases": evaluations,
    }
    output_path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def _percent(value: float | None) -> str:
    if value is None:
        return "не рассчитывалось"
    return f"{value * 100:.1f}%"


def _decimal(value: float | None) -> str:
    if value is None:
        return "не рассчитывалось"
    return f"{value:.2f}"


def _count_text(value: int | None) -> str:
    return str(value) if value is not None else "не рассчитывалось"


def _markdown_cell(value: object) -> str:
    return str(value).replace("|", "\\|").replace("\n", " ")


def _conclusion_lines(
    summary: dict[str, Any],
    grouped_metrics: dict[str, dict[str, Any]],
    evaluations: list[dict[str, Any]],
) -> list[str]:
    comparable_laws = {
        law: metrics
        for law, metrics in grouped_metrics.items()
        if metrics.get("top_5_accuracy") is not None
    }
    best_laws: list[str] = []
    weakest_laws: list[str] = []
    if comparable_laws:
        best_value = max(
            metrics["top_5_accuracy"]
            for metrics in comparable_laws.values()
        )
        weakest_value = min(
            metrics["top_5_accuracy"]
            for metrics in comparable_laws.values()
        )
        best_laws = [
            law
            for law, metrics in comparable_laws.items()
            if metrics["top_5_accuracy"] == best_value
        ]
        weakest_laws = [
            law
            for law, metrics in comparable_laws.items()
            if metrics["top_5_accuracy"] == weakest_value
        ]

    lines = [
        f"Проверено вопросов: {summary['questions_count']}.",
        (
            f"Recall@5 составляет {_percent(summary['top_5_accuracy'])}, "
            f"Recall@10 — {_percent(summary['top_10_accuracy'])}."
        ),
    ]
    if best_laws:
        lines.append(
            "Лучший результат по Recall@5: "
            + ", ".join(best_laws)
            + "."
        )
    if weakest_laws:
        lines.append(
            "Наибольшего внимания требует: "
            + ", ".join(weakest_laws)
            + "."
        )
    wrong_document_rate = summary.get("wrong_document_at_1")
    if wrong_document_rate:
        lines.append(
            "Есть подмешивание результатов из других законов на первой позиции: "
            f"Wrong Document@1 = {_percent(wrong_document_rate)}."
        )
    else:
        lines.append(
            "Проблема подмешивания других законов на первой позиции не выявлена."
        )
    return lines


def build_markdown_report(
    evaluations: list[dict[str, Any]],
    summary: dict[str, Any],
    grouped_metrics: dict[str, dict[str, Any]],
    metadata: dict[str, Any],
    question_type_metrics: dict[str, dict[str, Any]] | None = None,
) -> str:
    question_type_metrics = question_type_metrics or {}
    lines = [
        "# Отчёт о тестировании retrieval LawyerChat",
        "",
        f"- Дата и время запуска: {metadata['generated_at']}",
        f"- Количество оцениваемых вопросов: {summary['questions_count']}",
        f"- Out-of-scope кейсов: {summary['out_of_scope_count']}",
        f"- top_k: {metadata['top_k']}",
        f"- Retrieval mode: `{metadata['retrieval_mode']}`",
        f"- Модель embeddings: `{metadata['embedding_model']}`",
    ]
    if metadata["retrieval_mode"] == "hybrid":
        lines.append(
            (
                "- Hybrid weights: "
                f"semantic={metadata['hybrid_semantic_weight']}, "
                f"keyword={metadata['hybrid_keyword_weight']}, "
                f"metadata={metadata['hybrid_metadata_weight']}"
            )
        )
    lines.extend(
        [
            f"- База данных доступна: {metadata['database_available']}",
            f"- pgvector доступен: {metadata['pgvector_available']}",
            "",
            "## Итоговые метрики",
            "",
            "| Метрика | Значение |",
            "|---|---:|",
            f"| Recall@1 | {_percent(summary['top_1_accuracy'])} |",
            f"| Recall@3 | {_percent(summary['top_3_accuracy'])} |",
            f"| Recall@5 | {_percent(summary['top_5_accuracy'])} |",
            f"| Recall@10 | {_percent(summary['top_10_accuracy'])} |",
            f"| Recall@20 | {_percent(summary['top_20_accuracy'])} |",
            f"| MRR | {summary['mrr']:.4f} |",
            f"| Mean Rank | {_decimal(summary['mean_rank'])} |",
            (
                "| Document Recall@5 | "
                f"{_percent(summary['document_recall_at_5'])} |"
            ),
            (
                "| Wrong Document@1 | "
                f"{_percent(summary['wrong_document_at_1'])} |"
            ),
            "",
            "## Метрики по законам",
            "",
            (
                "| Закон | Вопросов | Recall@1 | Recall@5 | Recall@10 | "
                "Recall@20 | MRR | Mean Rank | Document Recall@5 | "
                "Wrong Document@1 |"
            ),
            "|---|---:|---:|---:|---:|---:|---:|---:|---:|---:|",
        ]
    )

    for law, metrics in grouped_metrics.items():
        lines.append(
            f"| {_markdown_cell(law)} | {metrics['questions_count']} | "
            f"{_percent(metrics['top_1_accuracy'])} | "
            f"{_percent(metrics['top_5_accuracy'])} | "
            f"{_percent(metrics['top_10_accuracy'])} | "
            f"{_percent(metrics['top_20_accuracy'])} | "
            f"{metrics['mrr']:.4f} | "
            f"{_decimal(metrics['mean_rank'])} | "
            f"{_percent(metrics['document_recall_at_5'])} | "
            f"{_percent(metrics['wrong_document_at_1'])} |"
        )

    lines.extend(
        [
            "",
            "## Метрики по типам вопросов",
            "",
            "| question_type | count | Recall@5 | Recall@10 | MRR |",
            "|---|---:|---:|---:|---:|",
        ]
    )
    for question_type, metrics in question_type_metrics.items():
        lines.append(
            f"| {_markdown_cell(question_type)} | "
            f"{metrics['questions_count']} | "
            f"{_percent(metrics['top_5_accuracy'])} | "
            f"{_percent(metrics['top_10_accuracy'])} | "
            f"{metrics['mrr']:.4f} |"
        )

    out_of_scope = [
        item
        for item in evaluations
        if item.get("not_evaluated")
    ]
    lines.extend(["", "## Out-of-scope кейсы", ""])
    if not out_of_scope:
        lines.append("Out-of-scope кейсы отсутствуют.")
    else:
        lines.append(
            "Эти вопросы перечислены отдельно и не входят в автоматические retrieval-метрики."
        )
        lines.append("")
        for evaluation in out_of_scope:
            case = evaluation["case"]
            lines.append(
                f"- `{case['id']}` — {case['question']} "
                f"(ожидаемое поведение: {case.get('expected_behavior', '—')})"
            )

    lines.extend(["", "## Ошибочные случаи", ""])
    failed = [
        item
        for item in evaluations
        if item.get("hit_top_5") is not True
    ]
    if not failed:
        lines.append("Правильная статья найдена в top-5 для всех вопросов.")
    else:
        for evaluation in failed:
            case = evaluation["case"]
            lines.extend(
                [
                    f"### {case['id']}",
                    "",
                    f"- Вопрос: {case['question']}",
                    f"- Тип вопроса: {case.get('question_type') or '—'}",
                    f"- Сложность: {case.get('difficulty') or '—'}",
                    (
                        "- Ожидаемые источники: "
                        + _format_expected_sources(case)
                    ),
                    (
                        "- Top-5 найденных результатов: "
                        + _format_top_results(
                            evaluation["top_results"],
                            limit=5,
                        )
                    ),
                    (
                        "- Совпадение: "
                        f"{evaluation.get('matched_relevance') or 'нет'}"
                    ),
                    f"- Комментарий: {case.get('comment') or '—'}",
                    "",
                ]
            )

    lines.extend(
        [
            "## Подробные результаты",
            "",
            (
                "| id | Закон | question_type | difficulty | Вопрос | "
                "Ожидаемые источники | rank | top-5 | Результат |"
            ),
            "|---|---|---|---|---|---|---:|---|---|",
        ]
    )
    for evaluation in evaluations:
        if evaluation.get("not_evaluated"):
            continue
        case = evaluation["case"]
        top_results = "; ".join(
            (
                f"{_result_document_name(result)} — ст. "
                f"{normalize_article_number(result.get('article_number')) or '—'}"
            )
            for result in evaluation["top_results"][:5]
        )
        lines.append(
            f"| {_markdown_cell(case['id'])} | {_markdown_cell(case.get('law', ''))} | "
            f"{_markdown_cell(case.get('question_type', ''))} | "
            f"{_markdown_cell(case.get('difficulty', ''))} | "
            f"{_markdown_cell(case.get('question', ''))} | "
            f"{_markdown_cell(_format_expected_sources(case))} | "
            f"{evaluation['rank'] or '—'} | {_markdown_cell(top_results)} | "
            f"{'успех' if evaluation.get('hit_top_5') else ('частично' if evaluation.get('partial_match') else 'ошибка')} |"
        )

    lines.extend(["", "## Вывод", ""])
    lines.extend(f"- {line}" for line in _conclusion_lines(
        summary,
        grouped_metrics,
        evaluations,
    ))
    return "\n".join(lines) + "\n"


def write_docx_report(
    evaluations: list[dict[str, Any]],
    summary: dict[str, Any],
    grouped_metrics: dict[str, dict[str, Any]],
    metadata: dict[str, Any],
    output_path: Path,
    question_type_metrics: dict[str, dict[str, Any]] | None = None,
) -> bool:
    question_type_metrics = question_type_metrics or {}
    try:
        from docx import Document
    except ImportError:
        print("WARNING: python-docx is not installed; DOCX report skipped", file=sys.stderr)
        return False

    document = Document()
    document.add_heading("Отчёт о тестировании retrieval LawyerChat", level=0)
    document.add_paragraph(f"Дата и время запуска: {metadata['generated_at']}")
    document.add_paragraph(
        f"Количество оцениваемых вопросов: {summary['questions_count']}"
    )
    document.add_paragraph(
        f"Out-of-scope кейсов: {summary['out_of_scope_count']}"
    )
    document.add_paragraph(f"top_k: {metadata['top_k']}")
    document.add_paragraph(f"Retrieval mode: {metadata['retrieval_mode']}")
    if metadata["retrieval_mode"] == "hybrid":
        document.add_paragraph(
            "Hybrid weights: "
            f"semantic={metadata['hybrid_semantic_weight']}, "
            f"keyword={metadata['hybrid_keyword_weight']}, "
            f"metadata={metadata['hybrid_metadata_weight']}"
        )
    document.add_paragraph(f"Модель embeddings: {metadata['embedding_model']}")
    document.add_paragraph(
        f"База данных: {metadata['database_available']}; "
        f"pgvector: {metadata['pgvector_available']}"
    )

    document.add_heading("Итоговые метрики", level=1)
    summary_table = document.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = "Метрика"
    summary_table.rows[0].cells[1].text = "Значение"
    for label, value in (
        ("Recall@1", _percent(summary["top_1_accuracy"])),
        ("Recall@3", _percent(summary["top_3_accuracy"])),
        ("Recall@5", _percent(summary["top_5_accuracy"])),
        ("Recall@10", _percent(summary["top_10_accuracy"])),
        ("Recall@20", _percent(summary["top_20_accuracy"])),
        ("MRR", f"{summary['mrr']:.4f}"),
        ("Mean Rank", _decimal(summary["mean_rank"])),
        ("Document Recall@5", _percent(summary["document_recall_at_5"])),
        ("Wrong Document@1", _percent(summary["wrong_document_at_1"])),
    ):
        cells = summary_table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    document.add_heading("Метрики по законам", level=1)
    law_table = document.add_table(rows=1, cols=6)
    law_table.style = "Table Grid"
    for cell, value in zip(
        law_table.rows[0].cells,
        (
            "Закон",
            "Вопросы",
            "Recall@1",
            "Recall@5",
            "Recall@10",
            "Recall@20",
        ),
    ):
        cell.text = value
    for law, metrics in grouped_metrics.items():
        cells = law_table.add_row().cells
        values = (
            law,
            str(metrics["questions_count"]),
            _percent(metrics["top_1_accuracy"]),
            _percent(metrics["top_5_accuracy"]),
            _percent(metrics["top_10_accuracy"]),
            _percent(metrics["top_20_accuracy"]),
        )
        for cell, value in zip(cells, values):
            cell.text = value

    document.add_paragraph("Дополнительные метрики по законам")
    additional_table = document.add_table(rows=1, cols=5)
    additional_table.style = "Table Grid"
    for cell, value in zip(
        additional_table.rows[0].cells,
        (
            "Закон",
            "MRR",
            "Mean Rank",
            "Document Recall@5",
            "Wrong Document@1",
        ),
    ):
        cell.text = value
    for law, metrics in grouped_metrics.items():
        cells = additional_table.add_row().cells
        values = (
            law,
            f"{metrics['mrr']:.4f}",
            _decimal(metrics["mean_rank"]),
            _percent(metrics["document_recall_at_5"]),
            _percent(metrics["wrong_document_at_1"]),
        )
        for cell, value in zip(cells, values):
            cell.text = value

    document.add_heading("Метрики по типам вопросов", level=1)
    type_table = document.add_table(rows=1, cols=5)
    type_table.style = "Table Grid"
    for cell, value in zip(
        type_table.rows[0].cells,
        ("question_type", "count", "Recall@5", "Recall@10", "MRR"),
    ):
        cell.text = value
    for question_type, metrics in question_type_metrics.items():
        cells = type_table.add_row().cells
        values = (
            question_type,
            str(metrics["questions_count"]),
            _percent(metrics["top_5_accuracy"]),
            _percent(metrics["top_10_accuracy"]),
            f"{metrics['mrr']:.4f}",
        )
        for cell, value in zip(cells, values):
            cell.text = value

    document.add_heading("Out-of-scope кейсы", level=1)
    out_of_scope = [
        item
        for item in evaluations
        if item.get("not_evaluated")
    ]
    if not out_of_scope:
        document.add_paragraph("Out-of-scope кейсы отсутствуют.")
    else:
        document.add_paragraph(
            "Эти вопросы не входят в автоматические retrieval-метрики."
        )
        for evaluation in out_of_scope:
            case = evaluation["case"]
            document.add_paragraph(
                f"{case['id']}: {case['question']} "
                f"(ожидаемое поведение: {case.get('expected_behavior', '—')})",
                style="List Bullet",
            )

    document.add_heading("Ошибочные случаи", level=1)
    failed = [
        item
        for item in evaluations
        if item.get("hit_top_5") is not True
    ]
    if not failed:
        document.add_paragraph(
            "Правильная статья найдена в top-5 для всех вопросов."
        )
    for evaluation in failed:
        case = evaluation["case"]
        document.add_heading(case["id"], level=2)
        document.add_paragraph(f"Вопрос: {case['question']}")
        document.add_paragraph(
            f"Тип вопроса: {case.get('question_type') or '—'}"
        )
        document.add_paragraph(
            f"Сложность: {case.get('difficulty') or '—'}"
        )
        document.add_paragraph(
            "Ожидаемые источники: "
            + _format_expected_sources(case)
        )
        document.add_paragraph(
            "Top-5 найденных результатов: "
            + _format_top_results(evaluation["top_results"], limit=5)
        )
        document.add_paragraph(
            "Совпадение: "
            + str(evaluation.get("matched_relevance") or "нет")
        )
        document.add_paragraph(f"Комментарий: {case.get('comment') or '—'}")

    document.add_heading("Подробные результаты", level=1)
    details_table = document.add_table(rows=1, cols=7)
    details_table.style = "Table Grid"
    for cell, value in zip(
        details_table.rows[0].cells,
        ("id", "Контекст", "Вопрос", "Ожидалось", "rank", "top-5", "Результат"),
    ):
        cell.text = value
    for evaluation in evaluations:
        if evaluation.get("not_evaluated"):
            continue
        case = evaluation["case"]
        top_articles = ", ".join(
            normalize_article_number(result.get("article_number")) or "—"
            for result in evaluation["top_results"][:5]
        )
        values = (
            case["id"],
            (
                f"{case.get('law') or ''}\n"
                f"question_type: {case.get('question_type') or '—'}\n"
                f"difficulty: {case.get('difficulty') or '—'}"
            ),
            case.get("question") or "",
            _format_expected_sources(case),
            str(evaluation["rank"] or "—"),
            top_articles,
            (
                "успех"
                if evaluation.get("hit_top_5")
                else (
                    "частично"
                    if evaluation.get("partial_match")
                    else "ошибка"
                )
            ),
        )
        cells = details_table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value

    document.add_heading("Вывод", level=1)
    for line in _conclusion_lines(summary, grouped_metrics, evaluations):
        document.add_paragraph(line, style="List Bullet")

    document.save(output_path)
    return True


def get_retrieval_search(retriever: Any, mode: str):
    if mode == "semantic":
        return retriever.search_semantic
    if mode == "hybrid":
        return retriever.search_hybrid
    raise ValueError(f"Unsupported retrieval mode: {mode}")


def run_evaluation(
    cases_path: Path,
    top_k: int,
    output_dir: Path,
    mode: str = "semantic",
) -> dict[str, Path]:
    from app.config import settings
    from app.db import (
        SessionLocal,
        check_database_connection,
        check_pgvector_extension,
    )
    from app.rag.embedder import Embedder
    from app.rag.retriever import Retriever

    if top_k < 1:
        raise RuntimeError("top_k must be greater than 0")

    cases = load_cases(cases_path)
    database_available = check_database_connection()
    pgvector_available = check_pgvector_extension()
    if not database_available:
        raise RuntimeError("Database is not available")

    db = SessionLocal()
    try:
        retriever = Retriever(db=db, embedder=Embedder())
        search = get_retrieval_search(retriever, mode)
        evaluations: list[dict[str, Any]] = []
        for index, case in enumerate(cases, start=1):
            print(f"[{index}/{len(cases)}] {case['id']}: {case['question']}")
            results = (
                []
                if case.get("case_type") == "out_of_scope"
                else search(case["question"], top_k=top_k)
            )
            evaluations.append(
                evaluate_case_results(case, results, top_k=top_k)
            )
    finally:
        db.close()

    summary = calculate_summary_metrics(evaluations)
    grouped_metrics = calculate_grouped_metrics(evaluations)
    question_type_metrics = calculate_question_type_metrics(evaluations)
    metadata = {
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "cases_count": len(cases),
        "top_k": top_k,
        "retrieval_mode": mode,
        "embedding_model": settings.embedding_model_name,
        "hybrid_semantic_weight": settings.hybrid_semantic_weight,
        "hybrid_keyword_weight": settings.hybrid_keyword_weight,
        "hybrid_metadata_weight": settings.hybrid_metadata_weight,
        "database_available": database_available,
        "pgvector_available": pgvector_available,
    }

    output_dir.mkdir(parents=True, exist_ok=True)
    paths = {
        "csv": output_dir / "retrieval_results.csv",
        "json": output_dir / "retrieval_results.json",
        "markdown": output_dir / "retrieval_report.md",
        "docx": output_dir / "retrieval_report.docx",
    }
    write_csv_report(evaluations, paths["csv"], metadata)
    write_json_report(
        evaluations,
        summary,
        grouped_metrics,
        metadata,
        paths["json"],
        question_type_metrics,
    )
    paths["markdown"].write_text(
        build_markdown_report(
            evaluations,
            summary,
            grouped_metrics,
            metadata,
            question_type_metrics,
        ),
        encoding="utf-8",
    )
    write_docx_report(
        evaluations,
        summary,
        grouped_metrics,
        metadata,
        paths["docx"],
        question_type_metrics,
    )
    return paths


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Evaluate LawyerChat retrieval on control questions."
    )
    parser.add_argument(
        "--cases",
        default="data/evaluation/retrieval_cases.json",
        help="Path to retrieval evaluation cases JSON",
    )
    parser.add_argument(
        "--top-k",
        type=int,
        default=20,
        help="Number of retrieval results per question",
    )
    parser.add_argument(
        "--mode",
        choices=("semantic", "hybrid"),
        default="semantic",
        help="Retrieval mode to evaluate",
    )
    parser.add_argument(
        "--output-dir",
        default="reports",
        help="Directory for CSV, JSON, Markdown and DOCX reports",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    try:
        paths = run_evaluation(
            cases_path=Path(args.cases),
            top_k=args.top_k,
            output_dir=Path(args.output_dir),
            mode=args.mode,
        )
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    print("Retrieval evaluation completed:")
    for report_type, path in paths.items():
        if path.exists():
            print(f"- {report_type}: {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
