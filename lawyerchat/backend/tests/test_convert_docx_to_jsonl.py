import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from scripts.convert_docx_to_jsonl import convert_file


class ConvertDocxToJsonlTests(unittest.TestCase):
    def test_docx_with_structure_and_articles_becomes_jsonl_records(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            input_path = Path(temp_dir) / "Семейный кодекс.docx"
            output_path = Path(temp_dir) / "semeynyy_kodeks_rf.jsonl"

            document = Document()
            document.add_paragraph("Раздел I. Общие положения")
            document.add_paragraph("Глава 1. Семейное законодательство")
            document.add_paragraph(
                "Статья 1. Основные начала семейного законодательства"
            )
            document.add_paragraph(
                "Семейное законодательство основывается на положениях статьи 2."
            )
            document.add_paragraph("Статья 2")
            document.add_paragraph("Отношения, регулируемые семейным законодательством.")
            document.save(input_path)

            records_count = convert_file(
                input_path,
                output_path,
                document_id="sk_rf",
                document_title="Семейный кодекс Российской Федерации",
            )
            records = [
                json.loads(line)
                for line in output_path.read_text(encoding="utf-8").splitlines()
            ]

        self.assertEqual(2, records_count)
        self.assertEqual(2, len(records))
        self.assertEqual(["1", "2"], [record["article_number"] for record in records])
        self.assertEqual(
            "Основные начала семейного законодательства",
            records[0]["article_title"],
        )
        self.assertIsNone(records[1]["article_title"])
        self.assertEqual("Раздел I. Общие положения", records[0]["section_title"])
        self.assertEqual(
            "Глава 1. Семейное законодательство",
            records[0]["chapter_title"],
        )
        self.assertEqual("docx", records[0]["source_format"])
        self.assertTrue(all(record["text"] for record in records))
        self.assertEqual(["2"], records[0]["referenced_articles"])


if __name__ == "__main__":
    unittest.main()
